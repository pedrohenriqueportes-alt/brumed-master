#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BRUMED MASTER — SaaS Web  v1.0
Portrex Operações e Soluções Integradas Ltda, 2026
"""

import io
import json
import os
import re
import statistics
import tempfile
import uuid
import zipfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import openpyxl
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

# ═══════════════════════════════════════════════════════════════════
#  CONSTANTES
# ═══════════════════════════════════════════════════════════════════

LIKERT = {
    "nunca": 1, "raramente": 2, "às vezes": 3,
    "as vezes": 3, "frequentemente": 4, "sempre": 5,
}
LIKERT_Q4 = {
    "nunca": 5, "raramente": 4, "às vezes": 3,
    "as vezes": 3, "frequentemente": 2, "sempre": 1,
}
BLOCKS = {
    "Q1": ("I",  "AH"),
    "Q2": ("AJ", "BL"),
    "Q3": ("BN", "CP"),
    "Q4": ("CR", "DG"),
    "Q5": ("DI", "DU"),
}
Q_LABELS = {
    "Q1": "Contexto de Trabalho",
    "Q2": "Exigências do Trabalho",
    "Q3": "Sintomas",
    "Q4": "Vivências Positivas",
    "Q5": "Vivências Negativas",
}

TEMPLATE_PATH = Path(__file__).parent / "resources" / "TEMPLATE OFICIAL.docx"
TEMP_DIR      = Path(tempfile.gettempdir()) / "brumed_saas"
TEMP_DIR.mkdir(exist_ok=True)


# ═══════════════════════════════════════════════════════════════════
#  LÓGICA DE NEGÓCIO  (preservada 100%)
# ═══════════════════════════════════════════════════════════════════

@dataclass
class Meta:
    razao_social: str
    cnpj: str
    profissional_responsavel: str
    ocupacao: str
    registro_profissional: str
    data_relatorio: str
    severidade_geral: Optional[int] = None
    severidades_por_setor: Optional[Dict[str, int]] = None


def normalize_text(value) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())

def normalize_key(value) -> str:
    return normalize_text(value).casefold()

def excel_col_to_int(col: str) -> int:
    result = 0
    for ch in col:
        result = result * 26 + (ord(ch.upper()) - 64)
    return result

def half_up_3(x: float) -> Decimal:
    return Decimal(str(x)).quantize(Decimal("0.000"), rounding=ROUND_HALF_UP)

def prob_from_mean(mean_3: Decimal) -> int:
    return int(mean_3.quantize(Decimal("1"), rounding=ROUND_HALF_UP))

def classify_brumed(prob: int, severidade: int) -> Tuple[str, str]:
    matrix = {
        5: {1:("PR2","RISCO ALTO"),2:("PR2","RISCO ALTO"),3:("PR2","RISCO ALTO"),4:("PR1","RISCO CRÍTICO"),5:("PR1","RISCO CRÍTICO")},
        4: {1:("PR3","RISCO MÉDIO"),2:("PR3","RISCO MÉDIO"),3:("PR2","RISCO ALTO"),4:("PR2","RISCO ALTO"),5:("PR1","RISCO CRÍTICO")},
        3: {1:("PR4","RISCO BAIXO"),2:("PR3","RISCO MÉDIO"),3:("PR3","RISCO MÉDIO"),4:("PR3","RISCO MÉDIO"),5:("PR2","RISCO ALTO")},
        2: {1:("PR4","RISCO BAIXO"),2:("PR4","RISCO BAIXO"),3:("PR4","RISCO BAIXO"),4:("PR3","RISCO MÉDIO"),5:("PR3","RISCO MÉDIO")},
        1: {1:("IRR","IRRELEVANTE"),2:("PR4","RISCO BAIXO"),3:("PR4","RISCO BAIXO"),4:("PR4","RISCO BAIXO"),5:("PR4","RISCO BAIXO")},
    }
    return matrix[prob][severidade]

def faixa_q3(media):
    m = float(media)
    if 1.00<=m<=1.79: return("manifestações pouco relatadas","Isso significa que, de forma geral, os trabalhadores não relatam presença relevante de sintomas físicos, psicológicos ou sociais associados à rotina de trabalho.")
    if 1.80<=m<=2.59: return("manifestações relatadas ocasionalmente","Isso significa que os sintomas são mencionados de forma pontual e não predominante, sugerindo presença esporádica de desconfortos ou percepções sintomáticas no contexto laboral.")
    if 2.60<=m<=3.39: return("manifestações relatadas com intermitência","Isso significa que há registro intermitente, porém não contínuo, de sintomas físicos, psicológicos ou sociais entre os trabalhadores, sugerindo atenção ao acompanhamento das condições organizacionais.")
    if 3.40<=m<=4.19: return("manifestações relatadas com regularidade","Isso significa que os sintomas são percebidos de forma regular entre os trabalhadores, indicando necessidade de monitoramento sistemático das condições relacionadas ao trabalho.")
    return("manifestações relatadas de forma recorrente","Isso significa que os sintomas são amplamente mencionados pelos trabalhadores, sugerindo presença significativa de repercussões percebidas no contexto ocupacional.")

def faixa_q4(media):
    m = float(media)
    if 1.00<=m<=1.79: return("forte presença de fatores protetivos","Isso significa que os trabalhadores percebem, de forma consistente, elementos positivos no ambiente organizacional, como reconhecimento, apoio e identificação com o trabalho.")
    if 1.80<=m<=2.59: return("boa presença de fatores protetivos","Isso significa que os trabalhadores relatam, de forma frequente, vivências positivas relacionadas ao ambiente de trabalho.")
    if 2.60<=m<=3.39: return("fatores protetivos moderados","Isso significa que as vivências positivas estão presentes, porém não de forma predominante entre os trabalhadores.")
    if 3.40<=m<=4.19: return("menor presença de fatores protetivos","Isso significa que as vivências positivas relacionadas ao ambiente de trabalho são percebidas de forma menos consistente pelos trabalhadores.")
    return("presença reduzida de fatores protetivos","Isso significa que os trabalhadores relatam baixa ocorrência de vivências positivas no contexto organizacional.")

def faixa_q5(media):
    m = float(media)
    if 1.00<=m<=1.79: return("vivências negativas pouco relatadas","Isso significa que os trabalhadores não relatam, de forma predominante, experiências emocionais negativas associadas ao contexto de trabalho.")
    if 1.80<=m<=2.59: return("vivências negativas relatadas ocasionalmente","Isso significa que experiências emocionais negativas são mencionadas de forma pontual pelos trabalhadores.")
    if 2.60<=m<=3.39: return("vivências negativas relatadas com intermitência","Isso significa que há percepção intermitente de experiências emocionais negativas relacionadas ao trabalho.")
    if 3.40<=m<=4.19: return("vivências negativas relatadas com regularidade","Isso significa que experiências emocionais negativas são percebidas de forma frequente pelos trabalhadores, sugerindo necessidade de acompanhamento das condições organizacionais.")
    return("vivências negativas relatadas de forma recorrente","Isso significa que há relato consistente de experiências emocionais negativas no contexto ocupacional, recomendando monitoramento contínuo das condições relacionadas ao trabalho.")

def convert_answer(raw, inverted=False):
    if raw is None or str(raw).strip() == "": return None
    if isinstance(raw, (int, float)):
        num = int(raw)
        if 1<=num<=5: return num if not inverted else {1:5,2:4,3:3,4:2,5:1}[num]
    key = normalize_key(raw)
    return (LIKERT_Q4 if inverted else LIKERT).get(key)

def replace_in_paragraph(paragraph, replacements):
    full = "".join(run.text for run in paragraph.runs)
    changed = False
    for old, new in replacements.items():
        if old in full:
            full = full.replace(old, new)
            changed = True
    if changed:
        for run in paragraph.runs: run.text = ""
        if paragraph.runs: paragraph.runs[0].text = full
        else: paragraph.add_run(full)

def iter_all_paragraphs(doc):
    for p in doc.paragraphs: yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: yield p

def analisar_respondente(row):
    respostas = []
    for q in ["Q1","Q2","Q3","Q4","Q5"]:
        respostas.extend([v for v in row[q] if v is not None])
    if not respostas:
        return {"total":0,"mais_comum":None,"pct_mais_comum":0,"desvio_padrao":0,
                "pct_extremos":0,"alertas":["Sem respostas válidas"],"incoerencias":[],
                "classificacao":"Sem dados suficientes",
                "medias_q":{q:0 for q in["Q1","Q2","Q3","Q4","Q5"]},"tem_inconsistencia":True}
    total=len(respostas); freq={}
    for r in respostas: freq[r]=freq.get(r,0)+1
    mais_comum=max(freq,key=freq.get); pct_mais_comum=round((freq[mais_comum]/total)*100,1)
    desvio=round(statistics.pstdev(respostas),3) if len(respostas)>1 else 0.0
    extremos=sum(1 for r in respostas if r in(1,5)); pct_extremos=round((extremos/total)*100,1)
    medias_q={q:round(sum(v for v in row[q] if v is not None)/len([v for v in row[q] if v is not None]),3) if [v for v in row[q] if v is not None] else 0 for q in["Q1","Q2","Q3","Q4","Q5"]}
    alertas,incoerencias=[],[]
    if pct_mais_comum>=85: alertas.append("Repetição muito elevada da mesma resposta")
    elif pct_mais_comum>=70: alertas.append("Repetição elevada da mesma resposta")
    if desvio<=0.45: alertas.append("Baixa variabilidade das respostas")
    if pct_extremos>=80: alertas.append("Predomínio de respostas extremas")
    elif pct_extremos>=65: alertas.append("Alta concentração de respostas extremas")
    if medias_q["Q4"]>=4.2 and medias_q["Q5"]>=4.2:
        incoerencias.append("Possível incoerência entre Q4 (vivências positivas) e Q5 (vivências negativas), ambos com médias elevadas.")
    if medias_q["Q3"]<=1.5 and medias_q["Q5"]>=4.5:
        incoerencias.append("Possível incoerência entre Q3 (sintomas muito baixos) e Q5 (vivências negativas muito altas).")
    if not alertas and not incoerencias: classificacao="Sem indícios relevantes de inconsistência"; tem=False
    elif len(alertas)+len(incoerencias)==1: classificacao="Atenção para consistência das respostas"; tem=True
    else: classificacao="Recomenda-se revisão técnica individual"; tem=True
    return {"total":total,"mais_comum":mais_comum,"pct_mais_comum":pct_mais_comum,
            "desvio_padrao":desvio,"pct_extremos":pct_extremos,"alertas":alertas,
            "incoerencias":incoerencias,"classificacao":classificacao,"medias_q":medias_q,"tem_inconsistencia":tem}


class BrumedProcessor:
    def __init__(self, xlsx_bytes: bytes, template_path: str, meta: Meta):
        self.xlsx_bytes    = xlsx_bytes
        self.template_path = template_path
        self.meta          = meta
        self.valid_rows:  List[dict]      = []
        self.setores:     List[str]       = []
        self.resultados:  Dict[str,dict]  = {}

    def load(self):
        self.wb = openpyxl.load_workbook(io.BytesIO(self.xlsx_bytes), data_only=True)
        if "Questionário" not in self.wb.sheetnames:
            raise ValueError("A planilha não contém a aba 'Questionário'.")
        self.ws = self.wb["Questionário"]

    def validate(self):
        if normalize_key(self.ws["D5"].value) not in ("nome setor","setor"):
            raise ValueError("Cabeçalho da coluna D inválido.")
        if normalize_key(self.ws["F5"].value) not in ("data resposta","data"):
            raise ValueError("Cabeçalho da coluna F inválido.")
        for q,(ini,fim) in BLOCKS.items():
            for c in range(excel_col_to_int(ini),excel_col_to_int(fim)+1):
                if normalize_text(self.ws.cell(row=5,column=c).value)=="":
                    raise ValueError(f"Cabeçalho vazio em {q}.")

    def collect_rows(self):
        for r in range(6,self.ws.max_row+1):
            if self.ws.cell(row=r,column=6).value in(None,""): continue
            setor=normalize_text(self.ws.cell(row=r,column=4).value)
            if not setor: continue
            row={"rownum":r,"setor":setor}; valid=True
            for q,(ini,fim) in BLOCKS.items():
                vals=[convert_answer(self.ws.cell(row=r,column=c).value,inverted=(q=="Q4"))
                      for c in range(excel_col_to_int(ini),excel_col_to_int(fim)+1)]
                if all(v is None for v in vals): valid=False
                row[q]=vals
            if valid: self.valid_rows.append(row)
        if not self.valid_rows: raise ValueError("Após o filtro F ≠ nulo, não restaram dados válidos.")
        self.setores=list(dict.fromkeys(r["setor"] for r in self.valid_rows))

    def _question_headers(self,q):
        ini,fim=BLOCKS[q]
        return[normalize_text(self.ws.cell(row=5,column=c).value)
               for c in range(excel_col_to_int(ini),excel_col_to_int(fim)+1)]

    def _mean(self,nums):
        vals=[n for n in nums if n is not None]
        return Decimal("0.000") if not vals else half_up_3(sum(vals)/len(vals))

    def _severidade_do_setor(self,setor):
        if self.meta.severidades_por_setor and setor in self.meta.severidades_por_setor:
            return self.meta.severidades_por_setor[setor]
        if self.meta.severidade_geral is not None: return self.meta.severidade_geral
        raise ValueError(f"Severidade não definida para: {setor}")

    def process(self):
        by_setor=defaultdict(list)
        for row in self.valid_rows: by_setor[row["setor"]].append(row)
        for setor,rows in by_setor.items():
            setor_res={}
            for q in["Q1","Q2","Q3","Q4","Q5"]:
                headers=self._question_headers(q); qm=[]; av=[]
                for idx,header in enumerate(headers):
                    vals=[r[q][idx] for r in rows if r[q][idx] is not None]
                    m=self._mean(vals); qm.append((header,m)); av.extend(vals)
                setor_res[q]={"media":self._mean(av),"top5":sorted(qm,key=lambda x:(-x[1],x[0]))[:5]}
            sev=self._severidade_do_setor(setor)
            for q in["Q1","Q2"]:
                prob=prob_from_mean(setor_res[q]["media"]); cod,cond=classify_brumed(prob,sev)
                setor_res[q].update({"prob":prob,"severidade":sev,"classif_codigo":cod,"classif_condicao":cond})
            self.resultados[setor]=setor_res

    def get_preview(self):
        """Retorna dados para preview na tela antes de gerar o Word."""
        preview={"setores":[],"total_respondentes":len(self.valid_rows)}
        RISCO_COLOR={"RISCO CRÍTICO":"#ef4444","RISCO ALTO":"#f97316",
                     "RISCO MÉDIO":"#eab308","RISCO BAIXO":"#22c55e","IRRELEVANTE":"#6b7280"}
        for setor in self.setores:
            r=self.resultados[setor]
            n=sum(1 for row in self.valid_rows if row["setor"]==setor)
            preview["setores"].append({
                "nome":setor,"respondentes":n,
                "q1":{"media":str(r["Q1"]["media"]).replace(".",","),
                      "codigo":r["Q1"]["classif_codigo"],
                      "condicao":r["Q1"]["classif_condicao"],
                      "cor":RISCO_COLOR.get(r["Q1"]["classif_condicao"],"#6b7280")},
                "q2":{"media":str(r["Q2"]["media"]).replace(".",","),
                      "codigo":r["Q2"]["classif_codigo"],
                      "condicao":r["Q2"]["classif_condicao"],
                      "cor":RISCO_COLOR.get(r["Q2"]["classif_condicao"],"#6b7280")},
                "q3":{"media":str(r["Q3"]["media"]).replace(".",",")},
                "q4":{"media":str(r["Q4"]["media"]).replace(".",",")},
                "q5":{"media":str(r["Q5"]["media"]).replace(".",",")},
            })
        return preview

    def cap6_lines(self):
        lines=[]
        for setor in self.setores:
            r=self.resultados[setor]; lines.append(("heading",setor))
            for q in["Q1","Q2","Q3","Q4","Q5"]:
                lines.append(("text",f"{q} – {Q_LABELS[q]}"))
                lines.append(("text",f"Média Geral: {str(r[q]['media']).replace('.',',')}"))
                if q in("Q1","Q2"):
                    lines.append(("text",f"Probabilidade Assumida: {r[q]['prob']}"))
                    lines.append(("text",f"Severidade Assumida: {r[q]['severidade']}"))
                    lines.append(("text",f"Classificação do Risco: {r[q]['classif_codigo']} – {r[q]['classif_condicao']}"))
            lines.append(("text","INTERPRETAÇÃO TÉCNICA SETORIAL"))
            q3f,q3c=faixa_q3(r["Q3"]["media"]); q4f,q4c=faixa_q4(r["Q4"]["media"]); q5f,q5c=faixa_q5(r["Q5"]["media"])
            for p in[
                f"O Questionário Q1 – Contexto de Trabalho apresentou classificação {r['Q1']['classif_codigo']} – {r['Q1']['classif_condicao']} conforme enquadramento na Matriz Oficial BRUMED.",
                f"O Questionário Q2 – Exigências do Trabalho apresentou classificação {r['Q2']['classif_codigo']} – {r['Q2']['classif_condicao']} conforme enquadramento na Matriz Oficial BRUMED.",
                f"O Questionário Q3 – Sintomas físico/psicológico/social apresentou média geral de {str(r['Q3']['media']).replace('.',',')}, indicando {q3f} no setor avaliado. {q3c}",
                f"O Questionário Q4 – Vivências positivas (escala invertida) apresentou média geral de {str(r['Q4']['media']).replace('.',',')}, indicando {q4f}. {q4c}",
                f"O Questionário Q5 – Vivências negativas apresentou média geral de {str(r['Q5']['media']).replace('.',',')}, indicando {q5f} no setor avaliado. {q5c}",
            ]: lines.append(("para",p))
            lines.append(("blank",""))
        return lines

    def cap7_rows(self):
        rows=[]
        for setor in self.setores:
            r=self.resultados[setor]
            rows.append([setor,str(r["Q1"]["media"]).replace(".",","),str(r["Q1"]["prob"]),
                         str(r["Q1"]["severidade"]),r["Q1"]["classif_codigo"],r["Q1"]["classif_condicao"],
                         str(r["Q2"]["media"]).replace(".",","),str(r["Q2"]["prob"]),
                         str(r["Q2"]["severidade"]),r["Q2"]["classif_codigo"],r["Q2"]["classif_condicao"]])
        return rows

    def render_docx(self,output_path:str):
        doc=Document(self.template_path)
        replacements={"<<RAZAO_SOCIAL>>":self.meta.razao_social,"<<CNPJ>>":self.meta.cnpj,
            "<<PROFISSIONAL_RESPONSAVEL>>":self.meta.profissional_responsavel,
            "<<OCUPACAO_PROFISSIONAL>>":self.meta.ocupacao,
            "<<NUMERO_DO_REGISTRO_PROFISSIONAL>>":self.meta.registro_profissional,
            "<<DATA_RELATORIO>>":self.meta.data_relatorio,
            "<<TOTAL_DE_RESPONDENTES_VALIDOS>>":str(len(self.valid_rows)),
            "<<DISTRIBUICAO_POR_SETOR>>":"; ".join(f"{s} ={sum(1 for r in self.valid_rows if r['setor']==s)}" for s in self.setores)}
        for p in iter_all_paragraphs(doc): replace_in_paragraph(p,replacements)
        cap6_anchor=cap7_anchor=None
        for p in doc.paragraphs:
            if "<<CAP6_CONTEUDO_COMPLETO>>" in p.text: cap6_anchor=p
            if "<<CAP7_LINHAS>>" in p.text: cap7_anchor=p
        if cap6_anchor is None or cap7_anchor is None:
            raise ValueError("Template sem placeholders de Capítulo 6 ou Capítulo 7.")
        cap6_anchor.text=""
        for kind,content in self.cap6_lines():
            p=cap6_anchor.insert_paragraph_before("")
            if kind=="heading": run=p.add_run(content); run.bold=True
            elif kind=="bullet": p.style="List Bullet"; p.add_run(content)
            else: p.add_run(content)
        cap7_anchor.text=""
        table=doc.add_table(rows=1,cols=11)
        for i,h in enumerate(["SETOR","Q1 MÉDIA","PROB Q1","SEVER Q1","CLASSIF Q1","CONDIÇÃO Q1",
                               "Q2 MÉDIA","PROB Q2","SEVER Q2","CLASSIF Q2","CONDIÇÃO Q2"]):
            table.rows[0].cells[i].text=h
        for row in self.cap7_rows():
            cells=table.add_row().cells
            for i,v in enumerate(row): cells[i].text=v
        body=doc._body._body; tbl=table._tbl; body.remove(tbl)
        idx=list(body).index(cap7_anchor._p)+1; body.insert(idx,tbl)
        doc.save(output_path)

    def render_top5(self,output_path:str):
        doc=Document()
        doc.add_heading("RELATÓRIO COMPLEMENTAR – TOP 5 FATORES POR QUESTIONÁRIO",level=1)
        doc.add_paragraph(f"Razão Social: {self.meta.razao_social}")
        doc.add_paragraph(f"CNPJ: {self.meta.cnpj}")
        doc.add_paragraph(f"Data do Relatório: {self.meta.data_relatorio}")
        for setor in self.setores:
            r=self.resultados[setor]; doc.add_heading(f"SETOR: {setor}",level=2)
            for q in["Q1","Q2","Q3","Q4","Q5"]:
                doc.add_heading(f"{q} – {Q_LABELS[q]}",level=3)
                for item,media in r[q]["top5"]:
                    doc.add_paragraph(f"{item} – {str(media).replace('.',',')}",style="List Number")
        doc.save(output_path)

    def render_consistencia(self,output_path:str):
        doc=Document()
        doc.add_heading("ANÁLISE DE CONSISTÊNCIA DAS RESPOSTAS",level=1)
        doc.add_paragraph(f"Razão Social: {self.meta.razao_social}")
        doc.add_paragraph(f"CNPJ: {self.meta.cnpj}")
        doc.add_paragraph(f"Data do Relatório: {self.meta.data_relatorio}")
        doc.add_paragraph(f"Respondentes válidos analisados: {len(self.valid_rows)}")
        doc.add_paragraph("Este relatório apresenta apenas os casos em que foram identificados indicadores de possível inconsistência nas respostas, com finalidade de apoio à revisão profissional. Os achados não representam, isoladamente, confirmação de má-fé, fraude ou invalidação automática do instrumento.")
        casos=0
        for idx,row in enumerate(self.valid_rows,start=1):
            analise=analisar_respondente(row)
            if not analise["tem_inconsistencia"]: continue
            casos+=1
            doc.add_heading(f"Respondente {idx} – Setor: {row['setor']}",level=2)
            doc.add_paragraph(f"Total de respostas válidas: {analise['total']}")
            doc.add_paragraph(f"Resposta mais frequente: {analise['mais_comum']}")
            doc.add_paragraph(f"Percentual da resposta mais frequente: {str(analise['pct_mais_comum']).replace('.',',')}%")
            doc.add_paragraph(f"Desvio padrão das respostas: {str(analise['desvio_padrao']).replace('.',',')}")
            doc.add_paragraph(f"Percentual de respostas extremas (1 e 5): {str(analise['pct_extremos']).replace('.',',')}%")
            doc.add_paragraph("Médias por questionário:")
            for q in["Q1","Q2","Q3","Q4","Q5"]:
                doc.add_paragraph(f"{q}: {str(analise['medias_q'][q]).replace('.',',')}",style="List Bullet")
            if analise["alertas"]:
                doc.add_paragraph("Indicadores gerais de consistência:")
                for a in analise["alertas"]: doc.add_paragraph(a,style="List Bullet")
            if analise["incoerencias"]:
                doc.add_paragraph("Incoerências entre questionários:")
                for i in analise["incoerencias"]: doc.add_paragraph(i,style="List Bullet")
            doc.add_paragraph(f"Classificação técnica: {analise['classificacao']}")
        if casos==0:
            doc.add_paragraph("Nenhum caso com indicador relevante de inconsistência foi identificado entre os respondentes válidos analisados.")
        doc.save(output_path)


# ═══════════════════════════════════════════════════════════════════
#  FASTAPI
# ═══════════════════════════════════════════════════════════════════

app = FastAPI(title="BRUMED MASTER SaaS", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"])


@app.get("/", response_class=HTMLResponse)
async def index():
    # Try static/index.html first, then fall back to root index.html
    html_path = Path(__file__).parent / "static" / "index.html"
    if not html_path.exists():
        html_path = Path(__file__).parent / "index.html"
    return HTMLResponse(html_path.read_text(encoding="utf-8"))


@app.post("/api/preview")
async def preview(
    arquivo: UploadFile = File(...),
    severidade_geral: int = Form(3),
):
    """Lê a planilha e retorna preview dos setores e respondentes."""
    try:
        xlsx_bytes = await arquivo.read()
        meta = Meta(razao_social="preview", cnpj="", profissional_responsavel="",
                    ocupacao="", registro_profissional="", data_relatorio="",
                    severidade_geral=severidade_geral)
        proc = BrumedProcessor(xlsx_bytes, str(TEMPLATE_PATH), meta)
        proc.load(); proc.validate(); proc.collect_rows(); proc.process()
        return JSONResponse({"ok": True, "data": proc.get_preview()})
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=400)


@app.post("/api/gerar")
async def gerar(
    arquivo: UploadFile = File(...),
    razao_social: str = Form(...),
    cnpj: str = Form(...),
    profissional_responsavel: str = Form(...),
    ocupacao: str = Form(...),
    registro_profissional: str = Form(...),
    data_relatorio: str = Form(...),
    severidade_geral: int = Form(...),
    severidades_por_setor: Optional[str] = Form(None),
    gerar_top5: bool = Form(True),
    gerar_consistencia: bool = Form(True),
):
    """Processa planilha e retorna ZIP com todos os relatórios."""
    try:
        xlsx_bytes = await arquivo.read()
        sev_setor_dict = None
        if severidades_por_setor:
            try:
                parsed = json.loads(severidades_por_setor)
                sev_setor_dict = {k: int(v) for k, v in parsed.items() if 1 <= int(v) <= 5}
            except Exception:
                sev_setor_dict = None

        meta = Meta(
            razao_social=razao_social, cnpj=cnpj,
            profissional_responsavel=profissional_responsavel,
            ocupacao=ocupacao, registro_profissional=registro_profissional,
            data_relatorio=data_relatorio, severidade_geral=severidade_geral,
            severidades_por_setor=sev_setor_dict,
        )

        if not TEMPLATE_PATH.exists():
            raise ValueError("Template não encontrado no servidor. Contacte o suporte.")

        proc = BrumedProcessor(xlsx_bytes, str(TEMPLATE_PATH), meta)
        proc.load(); proc.validate(); proc.collect_rows(); proc.process()

        # Gerar arquivos em pasta temporária
        session_id = str(uuid.uuid4())[:8]
        out_dir = TEMP_DIR / session_id
        out_dir.mkdir(exist_ok=True)

        empresa_safe = re.sub(r'[\\/*?:"<>|]', "", razao_social)[:40]
        data_safe    = data_relatorio.replace("/", "-")

        arquivos = []

        # Relatório principal
        path_principal = out_dir / f"RELATORIO_{empresa_safe}_{data_safe}.docx"
        proc.render_docx(str(path_principal))
        arquivos.append(path_principal)

        # Top 5
        if gerar_top5:
            path_top5 = out_dir / f"TOP5_FATORES_{empresa_safe}_{data_safe}.docx"
            proc.render_top5(str(path_top5))
            arquivos.append(path_top5)

        # Consistência
        if gerar_consistencia:
            path_cons = out_dir / f"CONSISTENCIA_{empresa_safe}_{data_safe}.docx"
            proc.render_consistencia(str(path_cons))
            arquivos.append(path_cons)

        # Compactar em ZIP
        zip_path = TEMP_DIR / f"BRUMED_{empresa_safe}_{session_id}.zip"
        with zipfile.ZipFile(zip_path, "w") as zf:
            for f in arquivos:
                zf.write(f, f.name)

        return FileResponse(
            path=str(zip_path),
            media_type="application/zip",
            filename=f"BRUMED_{empresa_safe}_{data_safe}.zip",
        )

    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=400)


@app.get("/health")
async def health():
    template_ok = TEMPLATE_PATH.exists()
    return {"status": "ok", "template": template_ok, "version": "1.0.0"}
